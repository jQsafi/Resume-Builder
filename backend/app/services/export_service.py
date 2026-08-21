import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.schemas.resume_schema import ResumeSchema

def generate_docx_resume(resume: ResumeSchema) -> bytes:
    doc = Document()
    
    # Page Margins (0.75 inch)
    sections = doc.sections
    for s in sections:
        s.top_margin = Inches(0.6)
        s.bottom_margin = Inches(0.6)
        s.left_margin = Inches(0.65)
        s.right_margin = Inches(0.65)

    primary_rgb = RGBColor(0, 104, 95)  # Deep Teal (#00685f)
    dark_rgb = RGBColor(25, 28, 30)     # Charcoal text

    # Name Header
    p_name = doc.add_paragraph()
    p_name.paragraph_format.space_after = Pt(2)
    p_name.paragraph_format.space_before = Pt(0)
    run_name = p_name.add_run(resume.personalInfo.fullName or "Your Name")
    run_name.font.size = Pt(22)
    run_name.font.bold = True
    run_name.font.name = "Arial"
    run_name.font.color.rgb = dark_rgb

    # Job Title
    p_role = doc.add_paragraph()
    p_role.paragraph_format.space_after = Pt(4)
    run_role = p_role.add_run(resume.personalInfo.jobTitle or "Professional Title")
    run_role.font.size = Pt(12)
    run_role.font.bold = True
    run_role.font.name = "Arial"
    run_role.font.color.rgb = primary_rgb

    # Contacts strip
    contacts = []
    if resume.personalInfo.email:
        contacts.append(resume.personalInfo.email)
    if resume.personalInfo.phone:
        contacts.append(resume.personalInfo.phone)
    if resume.personalInfo.location:
        contacts.append(resume.personalInfo.location)
    if resume.personalInfo.linkedin:
        contacts.append(resume.personalInfo.linkedin)
    if resume.personalInfo.github:
        contacts.append(resume.personalInfo.github)

    if contacts:
        p_contact = doc.add_paragraph()
        p_contact.paragraph_format.space_after = Pt(12)
        run_contact = p_contact.add_run("  •  ".join(contacts))
        run_contact.font.size = Pt(9.5)
        run_contact.font.color.rgb = RGBColor(100, 116, 139)

    # Section Helper
    def add_section_header(title: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(title.upper())
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = primary_rgb

    # Summary
    if resume.summary:
        add_section_header("Executive Summary")
        p_sum = doc.add_paragraph()
        p_sum.paragraph_format.space_after = Pt(10)
        run_sum = p_sum.add_run(resume.summary)
        run_sum.font.size = Pt(10)
        run_sum.font.color.rgb = dark_rgb

    # Technical Skills
    if resume.skills:
        add_section_header("Core Technical Competencies")
        if resume.skills.languages:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            r_label = p.add_run("Languages: ")
            r_label.font.bold = True
            r_label.font.size = Pt(9.5)
            r_val = p.add_run(", ".join(resume.skills.languages))
            r_val.font.size = Pt(9.5)

        if resume.skills.frameworks:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            r_label = p.add_run("Frameworks & DBs: ")
            r_label.font.bold = True
            r_label.font.size = Pt(9.5)
            r_val = p.add_run(", ".join(resume.skills.frameworks))
            r_val.font.size = Pt(9.5)

        if resume.skills.cloudDevops:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            r_label = p.add_run("Cloud & DevOps: ")
            r_label.font.bold = True
            r_label.font.size = Pt(9.5)
            r_val = p.add_run(", ".join(resume.skills.cloudDevops))
            r_val.font.size = Pt(9.5)

        if resume.skills.leadership:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            r_label = p.add_run("Architecture & Leadership: ")
            r_label.font.bold = True
            r_label.font.size = Pt(9.5)
            r_val = p.add_run(", ".join(resume.skills.leadership))
            r_val.font.size = Pt(9.5)

    # Experience
    if resume.experience:
        add_section_header("Professional Experience")
        for exp in resume.experience:
            p_exp = doc.add_paragraph()
            p_exp.paragraph_format.space_before = Pt(4)
            p_exp.paragraph_format.space_after = Pt(2)
            
            r_role = p_exp.add_run(exp.role or "Position")
            r_role.font.bold = True
            r_role.font.size = Pt(10.5)
            
            r_at = p_exp.add_run(" at ")
            r_at.font.size = Pt(10)
            
            r_comp = p_exp.add_run(exp.company or "Company")
            r_comp.font.bold = True
            r_comp.font.color.rgb = primary_rgb
            r_comp.font.size = Pt(10)
            
            date_str = f" ({exp.startDate} - {'Present' if exp.current else exp.endDate})"
            r_date = p_exp.add_run(date_str)
            r_date.font.size = Pt(9.5)
            r_date.font.italic = True
            r_date.font.color.rgb = RGBColor(100, 116, 139)

            for h in exp.highlights:
                p_bullet = doc.add_paragraph(style='List Bullet')
                p_bullet.paragraph_format.space_before = Pt(1)
                p_bullet.paragraph_format.space_after = Pt(2)
                r_b = p_bullet.add_run(h)
                r_b.font.size = Pt(9.5)
                r_b.font.color.rgb = dark_rgb

    # Education
    if resume.education:
        add_section_header("Education")
        for edu in resume.education:
            p_edu = doc.add_paragraph()
            p_edu.paragraph_format.space_before = Pt(3)
            p_edu.paragraph_format.space_after = Pt(2)
            
            r_deg = p_edu.add_run(edu.degree or "Degree")
            r_deg.font.bold = True
            r_deg.font.size = Pt(10)
            
            r_inst = p_edu.add_run(f", {edu.institution or 'Institution'}")
            r_inst.font.size = Pt(10)
            
            if edu.startDate or edu.endDate:
                r_dt = p_edu.add_run(f" ({edu.startDate} - {edu.endDate})")
                r_dt.font.size = Pt(9)
                r_dt.font.italic = True
                r_dt.font.color.rgb = RGBColor(100, 116, 139)

            if edu.details:
                p_det = doc.add_paragraph()
                p_det.paragraph_format.space_after = Pt(3)
                r_det = p_det.add_run(edu.details)
                r_det.font.size = Pt(9)
                r_det.font.color.rgb = RGBColor(100, 116, 139)

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
