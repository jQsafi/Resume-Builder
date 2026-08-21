import io
from docx import Document

def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extracts text from a .docx file buffer, including paragraphs and tables.
    """
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = []
    
    for p in doc.paragraphs:
        if p.text.strip():
            paragraphs.append(p.text.strip())
            
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                paragraphs.append(" | ".join(row_text))

    return "\n".join(paragraphs)
